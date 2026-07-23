import os
# from dotenv import load_dotenv
import requests
import json
import time
import fitz
import base64
import docx
import pandas as pd
import streamlit as st

# envpath = r"C:\Users\acehu\Downloads\AI automation\week 2 Python related Part\ai invoice pipeline project\api\.env"
# load_dotenv(envpath)


API_KEY = st.secrets["GEMINI_API_KEY"]


def pdf_extractor(file_path):
    pages = []
    pdf = fitz.open(file_path)
    for page in range(len(pdf)):

        pagedata = pdf.load_page(page)
        pic = pagedata.get_pixmap(dpi=150)

        # image_filename = "nothin to see here.png"

        inram = pic.tobytes("png")

        b64encodeddata = base64.b64encode(inram).decode("utf-8")

        pages.append(b64encodeddata)

    pdf.close()
    return pages


def docx_extractor(file_path):

    rawtext = docx.Document(file_path)
    database = ""
    for paragraph in rawtext.paragraphs:
        paragraph_text = paragraph.text
        database += paragraph_text + "\n"

    rawtext = docx.Document(file_path)
    tablestext = ""
    for table in rawtext.tables:
        for row in table.rows:
            database1 = []
            for cell in row.cells:
                database1.append(cell.text)
            tablestext += "\t".join(database1) + "\n"

    mastertext = database + "\n" + tablestext
    return mastertext


def image_extractor(file_path):
    with open(file_path, "rb") as image:
        b64encodeddata = base64.b64encode(image.read()).decode("utf-8")
    return [b64encodeddata]


def data_organizer(my_prompt, data, file_type):
    if file_type.endswith(".docx"):

        payload = {
            "contents": [
                {"parts": [{"text": my_prompt + "\n\n" + data + "\n\n" + file_type}]}
            ],
            "generationConfig": {
                "responseMimeType": "application/json",
                "temperature": 0.2,
            },
        }
    else:
        parts = [{"text": my_prompt + "\n\n" + file_type}]
        for image in data:
            parts.append({"inline_data": {"mime_type": "image/png", "data": image}})

        payload = {
            "contents": [{"parts": parts}],
            "generationConfig": {
                "responseMimeType": "application/json",
                "temperature": 0.2,
            },
        }

    headers = {
        "x-goog-api-key": API_KEY,
        "Content-Type": "application/json",
    }
    try:
        response = requests.post(
            "https://generativelanguage.googleapis.com/v1beta/models/gemini-3.1-flash-lite:generateContent",
            headers=headers,
            json=payload,
        )
        if response.status_code != 200:
            print(f"API Failed with status {response.status_code}: {response.text}")
            return None
        raw_text = response.json()

        our_required_text = raw_text["candidates"][0]["content"]["parts"][0]["text"]
        parsed_text = json.loads(our_required_text, strict=False)
        return parsed_text

    except Exception as e:
        print(
            f"this file was not processed due to some error. Try with it again {file_type} : {e}"
        )
        with open("failed files.txt", "a") as log:
            log.write(f"{file_type}: {e}\n")

        return None


def to_json(final_data, store_path):
    if final_data is None:
        print("No data to store. Skipping this entry.")
        return
    database = []
    if os.path.exists(store_path):
        print("File exists. Appending data to the existing file.")
        with open(store_path, "r") as file:
            database = json.load(file)

        if isinstance(final_data, list):
            database.extend(final_data)
        else:
            database.append(final_data)

        with open(store_path, "w") as file:
            json.dump(database, file, indent=4)

    else:
        print("File does not exist. Creating a new file and storing data.")
        if isinstance(final_data, list):
            database = final_data
        else:
            database = [final_data]
        with open(store_path, "w") as file:
            json.dump(database, file, indent=4)


def json_to_excel(json_file_path, excel_file_path):

    with open(json_file_path, "r") as file:
        json_text = json.load(file)

    clean_data = [item for item in json_text if item is not None]

    dataframe = pd.DataFrame(clean_data)

    # these are the results without cleaning the line items, so the line items are still in list format.

    # dataframe.to_csv(
    #     r"C:\Users\acehu\Downloads\AI automation\week 2 Python related Part\ai invoice pipeline project\excel file\result with out cleaning the line items.csv",
    #     index=False,
    # )

    #  these are the results with cleaning the line items, so the line items are in string format.
    dataframe["line_items"] = dataframe["line_items"].apply(
        lambda x: " , ".join(x) if (isinstance(x, list)) else x
    )
    print("Converting to Excel and CSV...")

    dataframe.to_excel(
        excel_file_path,
        index=False,
    )

    print("✅ Success! Check your folder for Final_Invoice_Report.xlsx")
    return None


def ai_invoice_pipeline(file_path, my_prompt, json_file_path, excel_file_path):
    files = os.listdir(file_path)
    print(f"these are the files in the directory: {files}" + ("\n"))
    for file in files:
        fullpath = os.path.join(file_path, file)
        print(f"Processing file: {file}" + ("\n"))

        if file.lower().endswith(".pdf"):
            print("PDF file detected" + ("\n"))
            pages = pdf_extractor(fullpath)

            final_data = data_organizer(my_prompt, pages, fullpath)
            to_json(final_data, json_file_path)

            continue

        elif file.lower().endswith(".docx"):
            print("DOCX file detected" + ("\n"))
            data = docx_extractor(fullpath)

        elif file.lower().endswith((".png", ".webp", ".jpg", ".jpeg")):
            print("Image file detected" + ("\n"))
            data = image_extractor(fullpath)

        final_data = data_organizer(my_prompt, data, fullpath)

        to_json(final_data, json_file_path)
    json_to_excel(json_file_path, excel_file_path)

    # with open(json_file_path, "rb") as jf:
    #     final_json_bytes = jf.read()

    # with open(excel_file_path, "rb") as ef:
    #     final_excel_bytes = ef.read()

    print("All files processed and data stored successfully.")

    # return final_json_bytes, final_excel_bytes
    return None


prompt = prompt = """
You are an expert financial data extraction AI. Extract the exact business information from the provided invoice/receipt document.

CRITICAL RULES:
1. Return ONLY a single, valid, raw JSON object. 
2. Do not wrap the output in markdown formatting (e.g., no ```json block). 
3. Do not include any conversational text before or after the JSON.
4. If a field is missing from the document, return null for that field (do not return "N/A" or "0").
5. All money amounts must be raw numbers (e.g., 150.50), do not include currency symbols in the number fields.

Use EXACTLY this JSON schema:
{
  "vendor_name": "Name of the company sending the invoice (From)",
  "vendor_address": "Full address and contact info of the vendor",
  "client_name": "Name of the person/company being billed (To / Bill To)",
  "client_address": "Full address of the client",
  "invoice_number": "The unique document identifier (e.g., INV-3337, 5475986675)",
  "invoice_date": "YYYY-MM-DD",
  "due_date": "YYYY-MM-DD",
  "subtotal_amount": 0.00,
  "tax_amount": 0.00,
  "shipping_amount": 0.00,
  "discount_amount": 0.00,
  "total_amount": 0.00,
  "currency": "USD, EUR, GBP, PKR, etc.",
  "payment_terms": "e.g., Net 30, Payable immediately",
  "line_items": ["Format as a single string per item. Example: '4x Hon Rocking Chair @ $461.48 = $1845.94'"],
  "is_receipt_or_invoice": "Receipt/Invoice/Quote",
  "file_name": "the name of the file being processed (e.g., invoice.pdf, receipt.jpg)"

}
"""

# filepath = r"C:\Users\acehu\Downloads\AI automation\week 2 Python related Part\ai invoice pipeline project\files"

# json_file_path = r"C:\Users\acehu\Downloads\AI automation\week 2 Python related Part\ai invoice pipeline project\result\Final_Invoice_Report.json"


# excel_file_path = r"C:\Users\acehu\Downloads\AI automation\week 2 Python related Part\ai invoice pipeline project\result\EXCEL\Final_Invoice_Report.xlsx"


# start = time.perf_counter()
# ai_invoice_pipeline(filepath, prompt, json_file_path, excel_file_path)
# end = time.perf_counter()

# print(f"Total run time: {end - start:.2f} seconds")
